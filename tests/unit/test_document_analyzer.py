"""
Test suite for Document Analyzer Agent.

Tests cover:
- PDF text extraction
- DOCX text extraction
- Unsupported file type handling
- FilePart byte extraction (base64 string and raw bytes)
- handler() — happy paths, edge cases, and error cases
- Agent integration (mocked)

Run with:
    pytest test_document_analyzer_agent.py -v
"""

import io
import base64
import pytest
import os
import sys
from unittest.mock import patch


# Helpers to generate minimal in-memory PDF and DOCX fixtures
def make_pdf_bytes(text: str = "Hello from PDF") -> bytes:
    """Create a minimal single-page PDF containing `text` using pypdf/reportlab."""
    try:
        from reportlab.pdfgen import canvas

        buf = io.BytesIO()
        c = canvas.Canvas(buf)
        c.drawString(100, 750, text)
        c.save()
        return buf.getvalue()
    except ImportError:
        # Fallback: minimal hand-crafted valid PDF
        content = (
            "%PDF-1.4\n"
            "1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
            "2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
            "3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]\n"
            "   /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\nendobj\n"
            "4 0 obj\n<< /Length 44 >>\nstream\n"
            "BT /F1 12 Tf 100 750 Td (" + text + ") Tj ET\nendstream\nendobj\n"
            "5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n"
            "xref\n0 6\n"
            "0000000000 65535 f \n"
            "trailer\n<< /Size 6 /Root 1 0 R >>\nstartxref\n9\n%%EOF"
        )
        return content.encode()


def make_docx_bytes(text: str = "Hello from DOCX") -> bytes:
    """Create a minimal DOCX file in memory containing `text`."""
    from docx import Document

    doc = Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def b64(data: bytes) -> str:
    return base64.b64encode(data).decode()


# Import the functions under test
# The example script lives outside of the package; add its directory to sys.path
# so imports succeed when running tests from the repo root.

sys.path.insert(
    0,
    os.path.abspath(
        os.path.join(
            os.path.dirname(__file__), "..", "..", "examples", "document-analyzer"
        )
    ),
)

from document_analyzer import (
    extract_text_from_pdf,
    extract_text_from_docx,
    extract_document_text,
    get_file_bytes,
    handler,
)


# PDF Extraction
class TestExtractTextFromPdf:
    def test_extracts_text_from_valid_pdf(self):
        pdf_bytes = make_pdf_bytes("Research methodology section")
        result = extract_text_from_pdf(pdf_bytes)
        assert isinstance(result, str)
        assert len(result) > 0

    def test_returns_string_for_empty_pdf(self):
        """A PDF with no text layers should return an empty string, not raise."""
        from pypdf import PdfWriter

        buf = io.BytesIO()
        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        writer.write(buf)
        result = extract_text_from_pdf(buf.getvalue())
        assert isinstance(result, str)

    def test_raises_on_invalid_bytes(self):
        with pytest.raises(Exception):
            extract_text_from_pdf(b"not a pdf at all")

    def test_multipage_pdf_joins_pages(self):
        try:
            from reportlab.pdfgen import canvas

            buf = io.BytesIO()
            c = canvas.Canvas(buf)
            c.drawString(100, 750, "Page one content")
            c.showPage()
            c.drawString(100, 750, "Page two content")
            c.save()
            result = extract_text_from_pdf(buf.getvalue())
            assert "Page one" in result
            assert "Page two" in result
        except ImportError:
            pytest.skip("reportlab not installed")


# DOCX Extraction
class TestExtractTextFromDocx:
    def test_extracts_text_from_valid_docx(self):
        docx_bytes = make_docx_bytes("Key findings of the study")
        result = extract_text_from_docx(docx_bytes)
        assert "Key findings" in result

    def test_returns_string_for_empty_docx(self):
        from docx import Document

        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        result = extract_text_from_docx(buf.getvalue())
        assert isinstance(result, str)

    def test_multiple_paragraphs_joined(self):
        from docx import Document

        doc = Document()
        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("Second paragraph.")
        buf = io.BytesIO()
        doc.save(buf)
        result = extract_text_from_docx(buf.getvalue())
        assert "First paragraph" in result
        assert "Second paragraph" in result

    def test_raises_on_invalid_bytes(self):
        with pytest.raises(Exception):
            extract_text_from_docx(b"not a docx file")


# extract_document_text dispatcher
class TestExtractDocumentText:
    def test_dispatches_pdf_mime(self):
        pdf_bytes = make_pdf_bytes("dispatch test pdf")
        result = extract_document_text(pdf_bytes, "application/pdf")
        assert isinstance(result, str)

    def test_dispatches_docx_mime(self):
        docx_bytes = make_docx_bytes("dispatch test docx")
        result = extract_document_text(
            docx_bytes,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        assert "dispatch test docx" in result

    def test_raises_for_unsupported_mime(self):
        with pytest.raises(ValueError, match="Unsupported file type"):
            extract_document_text(b"data", "image/png")

    def test_raises_for_empty_mime(self):
        with pytest.raises(ValueError, match="Unsupported file type"):
            extract_document_text(b"data", "")


# get_file_bytes
class TestGetFileBytes:
    def test_decodes_base64_string(self):
        raw = b"binary document content"
        file_part = {"file": {"bytes": b64(raw)}}
        assert get_file_bytes(file_part) == raw

    def test_returns_raw_bytes_directly(self):
        raw = b"already bytes"
        file_part = {"file": {"data": raw}}
        assert get_file_bytes(file_part) == raw

    def test_raises_when_no_data_key(self):
        file_part = {"file": {"uri": "http://example.com/file.pdf"}}
        with pytest.raises(ValueError, match="Unsupported file part format"):
            get_file_bytes(file_part)

    def test_raises_when_file_key_missing(self):
        with pytest.raises(KeyError):
            get_file_bytes({})


# handler() — integration tests (agent mocked)
MOCK_AGENT_RESPONSE = "This document discusses neural network training methodologies."


def make_file_part(file_bytes: bytes, mime: str, name: str = "test") -> dict:
    # include `text` to satisfy protocol validation; tests may not exercise
    # the endpoint layer but it's good to mirror real messages
    return {
        "kind": "file",
        "text": name,
        "file": {
            "bytes": b64(file_bytes),
            "mimeType": mime,
            "name": name,
        },
    }


def make_text_part(text: str) -> dict:
    return {"kind": "text", "text": text}


@pytest.fixture
def mock_agent():
    with patch("document_analyzer.agent") as m:
        m.run.return_value = MOCK_AGENT_RESPONSE
        yield m


class TestHandler:
    def test_returns_error_on_empty_messages(self, mock_agent):
        result = handler([])
        assert result == "No messages received."
        mock_agent.run.assert_not_called()

    def test_returns_error_when_no_file_in_messages(self, mock_agent):
        messages = [{"parts": [make_text_part("Summarize this document")]}]
        result = handler(messages)
        assert result == "No valid document found in the messages."
        mock_agent.run.assert_not_called()

    def test_handles_valid_pdf_message(self, mock_agent):
        pdf_bytes = make_pdf_bytes("Deep learning paper content")
        messages = [
            {
                "parts": [
                    make_text_part("What is the methodology?"),
                    make_file_part(pdf_bytes, "application/pdf"),
                ]
            }
        ]
        result = handler(messages)
        assert result == MOCK_AGENT_RESPONSE
        mock_agent.run.assert_called_once()
        call_input = mock_agent.run.call_args[1]["input"]
        assert "What is the methodology?" in call_input
        assert "Deep learning" in call_input

    def test_handles_valid_docx_message(self, mock_agent):
        docx_bytes = make_docx_bytes("Contract clause content")
        messages = [
            {
                "parts": [
                    make_text_part("Extract all obligations"),
                    make_file_part(
                        docx_bytes,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ]
            }
        ]
        result = handler(messages)
        assert result == MOCK_AGENT_RESPONSE
        call_input = mock_agent.run.call_args[1]["input"]
        assert "Extract all obligations" in call_input
        assert "Contract clause" in call_input

    def test_prompt_extracted_from_text_part(self, mock_agent):
        pdf_bytes = make_pdf_bytes("Some research content")
        messages = [
            {
                "parts": [
                    make_text_part("Identify the research gap"),
                    make_file_part(pdf_bytes, "application/pdf"),
                ]
            }
        ]
        handler(messages)
        call_input = mock_agent.run.call_args[1]["input"]
        assert "Identify the research gap" in call_input

    def test_multiple_files_combined(self, mock_agent):
        pdf_bytes = make_pdf_bytes("Document one content")
        docx_bytes = make_docx_bytes("Document two content")
        messages = [
            {
                "parts": [
                    make_text_part("Summarize both documents"),
                    make_file_part(pdf_bytes, "application/pdf"),
                    make_file_part(
                        docx_bytes,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ]
            }
        ]
        handler(messages)
        call_input = mock_agent.run.call_args[1]["input"]
        assert "Document one" in call_input
        assert "Document two" in call_input

    def test_unsupported_file_type_returns_error_in_doc(self, mock_agent):
        """Unsupported MIME should record an error string but not crash handler."""
        bad_file_part = {
            "kind": "file",
            "text": "bad.png",
            "file": {"bytes": b64(b"fake data"), "mimeType": "image/png", "name": "bad.png"},
        }
        messages = [{"parts": [make_text_part("analyze this"), bad_file_part]}]
        # The handler catches the exception and appends an error string
        result = handler(messages)
        # An error doc was appended, so agent still runs
        call_input = mock_agent.run.call_args[1]["input"]
        assert "Error processing file" in call_input

    def test_missing_parts_key_does_not_crash(self, mock_agent):
        """Messages without a 'parts' key should be skipped gracefully."""
        messages = [{"role": "user", "content": "no parts here"}]
        result = handler(messages)
        assert result == "No valid document found in the messages."

    def test_empty_prompt_still_runs_with_document(self, mock_agent):
        pdf_bytes = make_pdf_bytes("Research abstract")
        messages = [
            {
                "parts": [
                    make_file_part(pdf_bytes, "application/pdf"),
                    # No text part — prompt will be empty string
                ]
            }
        ]
        result = handler(messages)
        assert result == MOCK_AGENT_RESPONSE
        call_input = mock_agent.run.call_args[1]["input"]
        assert "Research abstract" in call_input

    def test_last_text_part_used_as_prompt(self, mock_agent):
        """If multiple text parts exist, the last one overwrites (current behavior)."""
        pdf_bytes = make_pdf_bytes("Some content")
        messages = [
            {
                "parts": [
                    make_text_part("First prompt"),
                    make_text_part("Final prompt"),
                    make_file_part(pdf_bytes, "application/pdf"),
                ]
            }
        ]
        handler(messages)
        call_input = mock_agent.run.call_args[1]["input"]
        assert "Final prompt" in call_input

    def test_multi_turn_conversation_uses_all_parts(self, mock_agent):
        """All messages in the conversation are iterated for files."""
        pdf_bytes = make_pdf_bytes("Uploaded in earlier turn")
        messages = [
            {
                "parts": [
                    make_file_part(pdf_bytes, "application/pdf"),
                ]
            },
            {
                "parts": [
                    make_text_part("Now summarize it"),
                ]
            },
        ]
        handler(messages)
        call_input = mock_agent.run.call_args[1]["input"]
        assert "Now summarize it" in call_input
        assert "Uploaded in earlier turn" in call_input
